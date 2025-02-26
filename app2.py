import streamlit as st
import pdfplumber
import re
import os
import tempfile
import pandas as pd
from datetime import datetime
from io import BytesIO
import base64

# Bibliotecas para DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Bibliotecas para fuzzy matching
from rapidfuzz import process, fuzz

# Bibliotecas para PDF (relatórios)
from fpdf import FPDF

###############################################################################
# CONFIGURAÇÕES E ESTADO
###############################################################################
st.set_page_config(page_title="Analista de Contracheques do INSS", layout="centered")

LOGO_PATH = "MP.png"  # Ajuste conforme o local do seu arquivo de logomarca
GLOSSARY_PATH = "Rubricas.txt"  # Ajuste conforme o local do seu arquivo de glossário

_fallback_state = {
    "df_informacoes": None,
    "df_descontos": None,
    "df_descontos_gloss": None,
    "df_descontos_gloss_sel": None,
    "nome_extraido": "",
    "nb_extraido": "",
    "valor_recebido": ""  # Fica vazio por padrão
}


def get_state_value(key):
    """Recupera um valor do estado."""
    try:
        return st.session_state[key]
    except:
        return _fallback_state.get(key, None)


def set_state_value(key, value):
    """Define um valor no estado."""
    try:
        st.session_state[key] = value
    except:
        _fallback_state[key] = value


###############################################################################
# FUNÇÕES AUXILIARES (LOGO, GLOSSÁRIO, FORMATOS)
###############################################################################
def get_image_base64(file_path):
    """Carrega uma imagem e retorna sua representação em base64."""
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def carregar_glossario(path):
    """Carrega o arquivo de glossário (Rubricas.txt) e retorna como lista de strings."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().splitlines()
    except Exception as e:
        st.error(f"Erro ao carregar glossário: {e}")
        return []


def inserir_totais_na_coluna(df, col_valor):
    """
    Insere linhas ao final da coluna col_valor com:
       - A = Valor Total (R$)
       - B = Valor Recebido – Autor (a)
       - Indébito (A-B)
       - Indébito em dobro (R$)

    *Utiliza o valor de 'valor_recebido' no estado para B.
    """
    if col_valor not in df.columns:
        return df

    def _to_float(x):
        try:
            return float(str(x).replace(',', '.').strip())
        except:
            return 0.0

    soma = df[col_valor].apply(_to_float).sum()
    if soma == 0:
        return df

    df_novo = df.copy()

    # Recupera o valor (string) e converte para float (para o cálculo)
    valor_recebido_str = get_state_value("valor_recebido") or "0"
    try:
        valor_recebido_num = float(str(valor_recebido_str).replace(",", ".").strip())
    except:
        valor_recebido_num = 0.0

    indebito = soma - valor_recebido_num
    indebito_dobro = 2 * indebito

    def en_us_format(number: float) -> str:
        return f"{number:,.2f}"

    # A = soma (formatado)
    A_str = en_us_format(soma)
    # B = valor recebido digitado (exatamente como string)
    B_str = valor_recebido_str
    indebito_str = en_us_format(indebito)
    indebito_dobro_str = en_us_format(indebito_dobro)

    # Insere as 4 linhas especiais no DataFrame
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [A_str], "DESCRIÇÃO": ["A = Valor Total (R$)"]})
    ], ignore_index=True)
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [B_str], "DESCRIÇÃO": ["B = Valor Recebido - Autor (a)"]})
    ], ignore_index=True)
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_str], "DESCRIÇÃO": ["Indébito (A-B)"]})
    ], ignore_index=True)
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_dobro_str], "DESCRIÇÃO": ["Indébito em dobro (R$)"]})
    ], ignore_index=True)

    # Limpa demais colunas nas linhas especiais
    mask_especial = df_novo["DESCRIÇÃO"].isin([
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ])
    for c in df_novo.columns:
        if c not in ["DESCRIÇÃO", col_valor]:
            df_novo.loc[mask_especial, c] = ""

    return df_novo


def formatar_valor_brl(valor):
    """Converte string no formato US '999.99' para '999,99'."""
    try:
        f = float(str(valor).replace(",", "").replace(".", "")) / 100
        return f"{f:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(valor)


def df_to_docx_bytes(dados: pd.DataFrame, titulo: str,
                     inserir_totais=False, col_valor_soma="DESCONTOS") -> bytes:
    """
    Converte DataFrame em um arquivo DOCX (bytes) com layout paisagem.
    Pode inserir linhas de total e demais itens se inserir_totais=True.
    """
    if inserir_totais:
        dados = inserir_totais_na_coluna(dados.copy(), col_valor_soma)

    # (3.1) Ajustar a numeração do NB (retirar vírgulas, substituir por pontos)
    # Exemplo: "137,939,448-9" => "137.939.448-9"
    import re
    titulo_fixed = titulo
    # Extrai a parte do NB, se existir, e troca vírgulas por ponto
    # Supondo que o título seja algo como: "Descontos Finais (Cronológico) - NOME - 137,939,448-9"
    match_nb = re.search(r"-(.*?)$", titulo)  # pega o final a partir do último hífen
    if match_nb:
        nb_dirty = match_nb.group(1)
        # Remove espaços extras
        nb_dirty_strip = nb_dirty.strip()
        # Substitui vírgulas por pontos
        nb_clean = nb_dirty_strip.replace(",", ".")
        # Constrói o novo título
        titulo_fixed = titulo.replace(nb_dirty_strip, nb_clean)

    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    titulo_heading = document.add_heading(titulo_fixed, level=1)
    titulo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if dados.empty:
        p = document.add_paragraph("DataFrame vazio - nenhum dado para exibir.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()

    colunas = dados.columns.tolist()
    table = document.add_table(rows=1, cols=len(colunas))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    for i, col_name in enumerate(colunas):
        hdr_cells[i].text = str(col_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    width_map = {}
    if "COD" in colunas:
        width_map["COD"] = 20
    if "DESCRIÇÃO" in colunas:
        width_map["DESCRIÇÃO"] = 130
    if "DESCONTOS" in colunas:
        width_map["DESCONTOS"] = 40
    if "DATA" in colunas:
        width_map["DATA"] = 30
    if "PÁGINA" in colunas:
        width_map["PÁGINA"] = 20

    for _, row in dados.iterrows():
        descricao = str(row.get("DESCRIÇÃO", ""))
        is_especial = descricao in [
            "A = Valor Total (R$)",
            "B = Valor Recebido - Autor (a)",
            "Indébito (A-B)",
            "Indébito em dobro (R$)"
        ]

        # (3.2) Ajustar casas decimais nas linhas especiais (exatamente como PDF).
        # Já foi feito na inserir_totais_na_coluna => ex: "1,608.90" => "1,608.90"
        # E depois iremos converter "." => "," em "ajustar_valores_docx" para exibição.
        row_cells = table.add_row().cells
        for i, col_name in enumerate(colunas):
            valor = str(row[col_name])
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run(valor)
            if col_name.upper() == "DESCRIÇÃO":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(9)
            if is_especial:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 0, 0)

        from docx.shared import Inches
        for i, col_name in enumerate(colunas):
            mm = width_map.get(col_name, 25)
            table.columns[i].width = Inches(mm / 25.4)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def ajustar_valores_docx(file_input_bytes: bytes) -> bytes:
    """
    Varre o DOCX gerado e converte qualquer valor 999.99 para 999,99
    (substituindo '.' por ',' no contexto de valores financeiros).
    """
    from docx import Document
    import re
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(file_input_bytes)
        tmp_in.flush()
        input_path = tmp_in.name

    output_path = input_path.replace(".docx", "_corrigido.docx")
    doc = Document(input_path)
    pattern = re.compile(r'([\d,]+\.\d{2})')
    for para in doc.paragraphs:
        found = pattern.findall(para.text)
        if found:
            for val_us in found:
                val_br = formatar_valor_brl(val_us)
                para.text = para.text.replace(val_us, val_br)
    doc.save(output_path)

    with open(output_path, "rb") as f:
        final_bytes = f.read()
    os.remove(input_path)
    os.remove(output_path)
    return final_bytes


def cruzar_descontos_com_rubricas(df_descontos, glossary, threshold=85):
    """
    Filtra linhas cujo texto em 'DESCRIÇÃO' combine (fuzzy matching)
    com itens do glossário acima de 'threshold' (0 a 100).
    """
    if df_descontos.empty or not glossary:
        return pd.DataFrame()

    unique_desc = df_descontos["DESCRIÇÃO"].unique()
    mapping = {}
    for desc in unique_desc:
        result = process.extractOne(desc, glossary, scorer=fuzz.ratio)
        mapping[desc] = (result is not None and result[1] >= threshold)

    mask = df_descontos["DESCRIÇÃO"].map(mapping)
    return df_descontos[mask]


###############################################################################
# FUNÇÕES COM pdfplumber (EXTRAÇÃO DE NOME, NB, COMPETÊNCIAS)
###############################################################################
def extrair_nome_e_nit_corrigido(pdf_path):
    """
    Extrai NB e Nome do PDF, a partir da primeira página.
    Exemplo de regex esperado:
      NB: 123.456.789-0
      Nome: JOAO DA SILVA
    Caso não encontre, retorna "N/D".
    """
    nome = "N/D"
    nb = "N/D"
    with pdfplumber.open(pdf_path) as pdf:
        if len(pdf.pages) > 0:
            text = pdf.pages[0].extract_text() or ""
            # Extração do NB (supondo que "NB:" ou "NIT:" => adaptado)
            nb_match = re.search(r"NB:\s*([\d\.\-]+)", text)
            if nb_match:
                nb = nb_match.group(1).strip()
            # Extração do Nome
            nome_match = re.search(r"Nome:\s*([A-Z\s]+)", text)
            if nome_match:
                nome = nome_match.group(1).strip().split("\n")[0]
    return nome, nb


def extrair_competencias_filtradas_por_contexto(pdf_path):
    """
    Extrai competências (MM/AAAA) se estiverem em linhas próximas
    à palavra "Competência" que aparece ao lado de "Período".
    """
    competencias_extraidas = []
    padrao_competencia = re.compile(r"\b(0[1-9]|1[0-2])\/(\d{4})\b")

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            linhas = text.split("\n")
            for i, linha in enumerate(linhas):
                if "Competência" in linha and "Período" in linha:
                    for j in range(1, 4):
                        if i + j < len(linhas):
                            matches = padrao_competencia.findall(linhas[i + j])
                            for (mes, ano) in matches:
                                competencias_extraidas.append(f"{mes}/{ano}")

    competencias_unicas = sorted(set(competencias_extraidas), key=lambda x: datetime.strptime(x, "%m/%Y"))
    df_competencias_filtradas = pd.DataFrame(competencias_unicas, columns=["Data Competência"])
    df_competencias_filtradas["Nome Competência"] = [f"Competência {i + 1}" for i in
                                                     range(len(df_competencias_filtradas))]
    return df_competencias_filtradas


def extrair_dados_contracheques_plumber(pdf_path):
    """
    Extrai dados essenciais do contracheque usando pdfplumber:
      - Código, Descrição Rubrica, Valor, Data (Competência), Página.
    Inicia extração após achar linha com:
      "Data de Início do Pagamento (DIP): dd/mm/aaaa MR: R$ <valores>"
    Ignora linhas com "Data de Nascimento".
    """
    dados_extracao = []
    iniciar_extracao = False
    padrao_DIP = re.compile(r"Data de Início do Pagamento \(DIP\): \d{2}/\d{2}/\d{4} MR: R\$ [\d.,]+")

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_number = page.page_number
            text = page.extract_text()
            if not text:
                continue
            linhas_filtradas = []

            for linha in text.split("\n"):
                if not iniciar_extracao:
                    if padrao_DIP.match(linha):
                        iniciar_extracao = True
                        continue
                    else:
                        continue

                if "Data de Nascimento" in linha:
                    continue

                linhas_filtradas.append(linha)

            competencia_match = re.search(r"Competência\s*(\d{2}/\d{4})", "\n".join(linhas_filtradas))
            competencia = competencia_match.group(1) if competencia_match else "N/A"

            rubrica_detectada = False
            for linha in linhas_filtradas:
                if "RUBRICA" in linha.upper():
                    dados_extracao.append({
                        "Código": "Rubrica",
                        "Descrição Rubrica": "Descrição Rubrica",
                        "Valor": "Valor",
                        "Data": competencia,
                        "Página": page_number
                    })
                    rubrica_detectada = True
                    continue

                parts = linha.split()
                if len(parts) >= 3 and parts[0].isdigit():
                    if rubrica_detectada:
                        rubrica_detectada = False
                    codigo = parts[0]
                    descricao = " ".join(parts[1:-1]).replace("R$", "").strip()
                    valor = parts[-1]
                    dados_extracao.append({
                        "Código": codigo,
                        "Descrição Rubrica": descricao,
                        "Valor": valor,
                        "Data": competencia,
                        "Página": page_number
                    })

    df = pd.DataFrame(dados_extracao)
    current_segment = 0
    intervalos = []
    for _, row in df.iterrows():
        if row["Código"] == "Rubrica":
            current_segment += 1
            intervalos.append("")
        else:
            intervalos.append(f"Competência {current_segment}" if current_segment > 0 else "")

    df["Intervalos"] = intervalos
    return df


def criar_informacoes_com_datas(df_rubricas, df_competencias):
    """
    Associa cada linha do DataFrame de Rubricas às datas extraídas (df_competencias),
    de acordo com a coluna 'Intervalos' => "Competência X".
    """
    df_info = df_rubricas.copy()
    for idx, row in df_info.iterrows():
        intervalo = row.get("Intervalos", "")
        match = re.search(r"Competência\s+(\d+)", intervalo)
        if match:
            num_comp = int(match.group(1))
            if 0 <= num_comp - 1 < len(df_competencias):
                data_comp = df_competencias.loc[num_comp - 1, "Data Competência"]
                df_info.at[idx, "Data"] = data_comp
    return df_info


###############################################################################
# CLASSE PDFBASICO PARA O RELATÓRIO BÁSICO
###############################################################################
class PDFBasico(FPDF):
    """
    Ajusta cabeçalho do relatório, incluindo 'Contracheque ISS - nome + NB'.
    """

    def __init__(self, nome_user, nb_user, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.nome_user = nome_user
        self.nb_user = nb_user

    def header(self):
        self.set_font('Arial', 'B', 12)
        titulo = f"Contracheque ISS - {self.nome_user} - {self.nb_user}"
        self.cell(0, 10, titulo, border=False, ln=True, align='C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Página {self.page_no()}', border=False, ln=False, align='C')


def salvar_em_pdf_basico(dados, file_name, nome_user, nb_user):
    """
    Gera um PDF simples com colunas: ["Código", "Descrição Rubrica", "Valor", "Data", "Página"].
    Usa a classe PDFBasico com cabeçalho personalizado (nome + NB).
    """
    headers = ["Código", "Descrição Rubrica", "Valor", "Data", "Página"]
    col_widths = {
        "Código": 30,
        "Descrição Rubrica": 130,
        "Valor": 40,
        "Data": 40,
        "Página": 40
    }

    pdf = PDFBasico(nome_user=nome_user, nb_user=nb_user, orientation='L', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.set_fill_color(200, 220, 255)

    # Cabeçalho das colunas
    for title in headers:
        pdf.cell(col_widths[title], 10, title, border=1, align='C', fill=True)
    pdf.ln()

    # Dados
    for _, row in dados.iterrows():
        for col in headers:
            text = str(row.get(col, ""))
            pdf.cell(col_widths[col], 10, text, border=1, align='C')
        pdf.ln()

    pdf.output(file_name)


###############################################################################
# MAIN
###############################################################################
def main():
    # Exibir logomarca
    logo_base64 = get_image_base64(LOGO_PATH)
    if logo_base64:
        st.markdown(f"""
            <div style="text-align: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{logo_base64}" alt="Logomarca" style="width: 300px;">
            </div>
            """, unsafe_allow_html=True)

    # Título principal
    st.title("Analista de Contracheques do INSS")

    # Upload do PDF
    uploaded_file = st.file_uploader(
        "Clique no botão para enviar o arquivo PDF (Contracheque INSS)",
        type="pdf"
    )

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        try:
            # Extrair Nome e NB
            nome_final, nb_final = extrair_nome_e_nit_corrigido(tmp_file_path)
            set_state_value("nome_extraido", nome_final)
            set_state_value("nb_extraido", nb_final)

            # Extrair competências
            df_competencias = extrair_competencias_filtradas_por_contexto(tmp_file_path)
            # Extrair dados do contracheque
            df_final = extrair_dados_contracheques_plumber(tmp_file_path)
            if df_final is None or df_final.empty:
                st.warning("Não foram encontradas informações no PDF.")
                return

            # Associa rubricas às datas (competências)
            df_informacoes = criar_informacoes_com_datas(df_final, df_competencias)
            df_informacoes = df_informacoes[df_informacoes["Código"] != "Rubrica"]
            if "Intervalos" in df_informacoes.columns:
                df_informacoes.drop(columns=["Intervalos"], inplace=True)

            st.subheader("Informações extraídas com datas")
            st.dataframe(df_informacoes, use_container_width=True)

            nome_user = get_state_value("nome_extraido") or "ND"
            nb_user = get_state_value("nb_extraido") or "ND"

            # (1) Retirar o texto "ISS_" do nome do PDF
            base_pdf_name = f"Contracheque {nome_user}_{nb_user}.pdf"
            pdf_info_path = os.path.join(tempfile.gettempdir(), base_pdf_name)
            salvar_em_pdf_basico(df_informacoes, pdf_info_path, nome_user, nb_user)

            with open(pdf_info_path, "rb") as pdf_file:
                st.download_button(
                    "Baixar Informações com Datas (PDF)",
                    data=pdf_file,
                    file_name=base_pdf_name,
                    mime="application/pdf"
                )

            set_state_value("df_informacoes", df_informacoes)

        finally:
            os.unlink(tmp_file_path)

    # Recupera DataFrame principal
    df_informacoes = get_state_value("df_informacoes")
    nome_user = get_state_value("nome_extraido") or "ND"
    nb_user = get_state_value("nb_extraido") or "ND"

    if df_informacoes is not None and not df_informacoes.empty:
        # Lista de Rubricas
        st.markdown("## Lista de Rúbricas")
        glossary_terms = carregar_glossario(GLOSSARY_PATH)
        if glossary_terms:
            df_rubricas = pd.DataFrame({"Rubricas": glossary_terms})
            st.dataframe(df_rubricas, use_container_width=True)
        else:
            st.warning("Glossário vazio ou não encontrado.")

        # Ajustar colunas para "DESCRIÇÃO", "DESCONTOS", "PÁGINA"
        df_aux = df_informacoes.copy()
        df_aux.rename(columns={
            "Descrição Rubrica": "DESCRIÇÃO",
            "Valor": "DESCONTOS",
            "Página": "PÁGINA"
        }, inplace=True)

        # Filtrar Descontos no Glossário
        st.markdown("## Filtrar Descontos no Glossário")
        with st.form("form_filtro_gloss"):
            thresh = st.slider("Nível de Similaridade (0.1 a 1.0)", 0.1, 1.0, 0.85, 0.1)
            submit_gloss = st.form_submit_button("Filtrar com Rubricas")

        if submit_gloss:
            threshold_value = int(thresh * 100)
            df_descontos = df_aux[df_aux["DESCONTOS"].str.strip() != ""].copy()
            set_state_value("df_descontos", df_descontos)
            df_desc_gloss = cruzar_descontos_com_rubricas(df_descontos, glossary_terms, threshold_value)
            set_state_value("df_descontos_gloss", df_desc_gloss)
            set_state_value("df_descontos_gloss_sel", None)

        df_descontos_gloss = get_state_value("df_descontos_gloss")
        if df_descontos_gloss is not None and not df_descontos_gloss.empty:
            st.markdown("### Descontos x Glossário")
            st.dataframe(df_descontos_gloss, use_container_width=True)

            st.markdown("## Lista única de descontos")
            df_sel = get_state_value("df_descontos_gloss_sel")
            if df_sel is None:
                df_sel = df_descontos_gloss

            with st.form("form_inclusao_descontos"):
                valores_unicos = sorted(df_sel["DESCRIÇÃO"].unique())
                st.write("Marque os itens que deseja incluir:")
                selected_descr = []
                for i, val in enumerate(valores_unicos):
                    qtd = df_sel[df_sel["DESCRIÇÃO"] == val].shape[0]
                    label_str = f"{i + 1} - {val} (Qtd: {qtd})"
                    if st.checkbox(label_str, key=f"chk_{i}"):
                        selected_descr.append(val)
                incluir_btn = st.form_submit_button("Confirmar Inclusão (Descontos)")

            if incluir_btn:
                if selected_descr:
                    df_incluido = df_sel[df_sel["DESCRIÇÃO"].isin(selected_descr)].copy()
                    set_state_value("df_descontos_gloss_sel", df_incluido)
                    st.success("Descontos selecionados com sucesso!")
                    st.markdown("### Lista restantes após exclusões")
                    st.dataframe(df_incluido, use_container_width=True)
                else:
                    st.warning("Nenhuma descrição selecionada.")

            df_final_sel = get_state_value("df_descontos_gloss_sel")
            if df_final_sel is not None and not df_final_sel.empty:
                st.markdown("## Apresentar Rúbricas para Débitos (Descontos Finais)")

                df_final = df_final_sel.copy()
                # Ajusta páginas (caso esteja em branco)
                df_final["PÁGINA"] = pd.to_numeric(df_final["PÁGINA"], errors='coerce').fillna(0)
                # Ordena por Data + Página
                df_final = df_final.sort_values(
                    by=["Data", "PÁGINA"],
                    key=lambda col: pd.to_datetime(col, format="%m/%Y", errors='coerce')
                ).reset_index(drop=True)

                # Apenas colunas relevantes
                df_final = df_final[["Código", "DESCRIÇÃO", "DESCONTOS", "Data"]]

                # Cálculo de A (sem apresentar botão)
                def _to_float(x):
                    try:
                        return float(str(x).replace(',', '.').strip())
                    except:
                        return 0.0

                A_val = df_final["DESCONTOS"].apply(_to_float).sum()
                A_str = f"{A_val:,.2f}"

                # Exibe A diretamente em tela (sem botão)
                st.write(f"A = Valor Total (R$): {A_str}")

                col1, col2 = st.columns(2)
                with col1:
                    # Input do valor B
                    valor_recebido_input = st.text_input("B = Valor Recebido - Autor (a)", "0")

                try:
                    vrnum = float(valor_recebido_input.replace(',', '.').strip())
                except:
                    vrnum = 0.0

                indebito_val = A_val - vrnum
                indebito_dobro_val = 2 * indebito_val
                indebito_str = f"{indebito_val:,.2f}"
                indebito_dobro_str = f"{indebito_dobro_val:,.2f}"

                with col2:
                    st.write(f"Indébito (A-B): {indebito_str}")
                    st.write(f"Indébito em dobro (R$): {indebito_dobro_str}")

                # Armazena o valor digitado no estado
                set_state_value("valor_recebido", valor_recebido_input)

                with st.form("form_descontos_finais"):
                    submit_final = st.form_submit_button("Gerar Relatório Final com Descontos")

                if submit_final:
                    # Monta título final (3.1 => trocar vírgula por ponto no NB)
                    nb_user_fixed = nb_user.replace(",", ".")
                    titulo_final = f"Descontos Finais (Cronológico) - {nome_user} - {nb_user_fixed}"

                    df_com_totais = inserir_totais_na_coluna(df_final.copy(), "DESCONTOS")

                    # (1) Retirar "ISS_" do nome do PDF
                    pdf_final_name = f"Contracheque Descontos_Finais_{nome_user}_{nb_user}.pdf"
                    pdf_final_path = os.path.join(tempfile.gettempdir(), pdf_final_name)

                    # Geração do PDF final
                    pdf = FPDF(orientation='L', format='A4')
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(0, 10, titulo_final, border=False, ln=True, align='C')
                    pdf.ln(10)

                    pdf.set_font("Arial", "B", 10)
                    headers = ["Código", "DESCRIÇÃO", "DESCONTOS", "Data"]
                    col_widths = {
                        "Código": 25,
                        "DESCRIÇÃO": 150,
                        "DESCONTOS": 35,
                        "Data": 40
                    }
                    pdf.set_fill_color(200, 220, 255)
                    for h in headers:
                        pdf.cell(col_widths[h], 8, h, border=1, align='C', fill=True)
                    pdf.ln()

                    # Monta as linhas no PDF
                    for _, row in df_com_totais.iterrows():
                        desc = row["DESCRIÇÃO"]
                        is_especial = desc in [
                            "A = Valor Total (R$)",
                            "B = Valor Recebido - Autor (a)",
                            "Indébito (A-B)",
                            "Indébito em dobro (R$)"
                        ]
                        if is_especial:
                            # (2.1) Fonte maior
                            pdf.set_font("Arial", "B", 12)
                            # (2.2) Cor vermelha
                            pdf.set_text_color(255, 0, 0)
                        else:
                            pdf.set_font("Arial", "", 9)
                            pdf.set_text_color(0, 0, 0)

                        row_data = []
                        for h in headers:
                            val = str(row[h])
                            if h == "DESCONTOS" and val.strip():
                                # Converte para formato BRL
                                val = formatar_valor_brl(val)
                            row_data.append(val)

                        for h, val in zip(headers, row_data):
                            pdf.cell(col_widths[h], 8, val, border=1, align='C')
                        pdf.ln()

                    pdf.output(pdf_final_path)

                    with open(pdf_final_path, "rb") as f_pdf:
                        st.download_button(
                            "Baixar PDF (Descontos Finais)",
                            data=f_pdf.read(),
                            file_name=pdf_final_name,
                            mime="application/pdf"
                        )

                    # Geração do DOCX final
                    # (1) Retirar "ISS_" do nome do DOCX
                    docx_final_name = f"Contracheque Descontos_Finais_{nome_user}_{nb_user}.docx"

                    docx_bytes = df_to_docx_bytes(
                        dados=df_final.copy(),
                        titulo=titulo_final,
                        inserir_totais=True,
                        col_valor_soma="DESCONTOS"
                    )
                    docx_bytes_corrigido = ajustar_valores_docx(docx_bytes)

                    st.download_button(
                        label="Baixar DOCX (Descontos Finais)",
                        data=docx_bytes_corrigido,
                        file_name=docx_final_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


if __name__ == "__main__":
    main()

